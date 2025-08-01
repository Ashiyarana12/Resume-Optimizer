from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
from io import BytesIO
import re

def create_template_resume_docx(text: str) -> str:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    section_titles = [
        "Name + Contact Info", "Summary", "Education",
        "Experience", "Projects", "Skills"
    ]

    # Split the text based on numbered sections (e.g., "1. ", "2. ")
    sections = re.split(r"\n?\d+\.\s+", text.strip())

    # --- ADDED FOR DEBUGGING ---
    print("--- Parsed Sections from LLM Output Start ---")
    # The first element after split will often be an empty string if the text starts with a number,
    # or the content before the first number. We're primarily interested in the numbered sections.
    for i, sec in enumerate(sections):
        print(f"Section {i}: \n{sec.strip()}\n---")
    print("--- Parsed Sections from LLM Output End ---")
    # ---------------------------

    # The first element of 'sections' after splitting by '1. ' might be empty or pre-content.
    # We expect the actual "Name + Contact Info" to be in sections[1] if the split worked as intended.
    if len(sections) > 1:
        name_block = sections[1].strip()
        name_para = doc.add_paragraph(name_block)
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in name_para.runs:
            run.font.size = Pt(14)
            run.bold = True
        doc.add_paragraph() # Add a blank line after name/contact

    # Iterate through the rest of the sections, starting from the "Summary" (index 2 in sections list)
    # and mapping them to predefined titles.
    # Note: section_titles are 0-indexed, and sections list might have an initial empty string or pre-header content.
    # So, sections[2] should correspond to section_titles[1] (Summary).
    for i in range(2, len(sections)):
        # Ensure we don't go out of bounds for section_titles
        if (i - 1) < len(section_titles):
            title = section_titles[i - 1]
            content = sections[i].strip()

            # Add section heading
            heading_para = doc.add_paragraph()
            run = heading_para.add_run(title)
            run.bold = True
            run.underline = True
            run.font.size = Pt(12)
            doc.add_paragraph() # Add a blank line after heading

            # Add content lines, handling bullet points
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue # Skip empty lines

                # Check for common bullet point indicators
                if line.startswith("-") or line.startswith("•"):
                    doc.add_paragraph(line, style="List Bullet")
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(4) # Small space after each line

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    encoded = base64.b64encode(buffer.read()).decode("utf-8")
    return encoded